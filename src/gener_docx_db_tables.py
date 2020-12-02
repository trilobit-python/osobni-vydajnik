#!/usr/bin/env python
# -*- coding: windows-1250 -*-
# -------------------------------------------------------------------------------
# Name:        Generuje z Oracle schematu - popis tabulek do docx dokumentu
# Author:      P3402617
# Created:     16.10.2020
# Copyright:   NemecDa
# -------------------------------------------------------------------------------

import tkinter as tk

import openpyxl
import pandas as pd
from tkinter import filedialog

from xlsxwriter import Workbook

from src.utils.oracle_database import OracleDatabase


def doctable(data, tabletitle, pathfile):
    from docx import Document
    import pandas as pd
    document = Document(pathfile)

    #
    for p in document.paragraphs:
        if 'Seznam nových tabulek' in p.text:
            p.clear()
            df = pd.DataFrame(data)  # My input data is in the 2D list form

            # add a table to the end and create a reference variable
            # extra row is so we can add the header row
            t = document.add_table(df.shape[0] + 1, df.shape[1])

            # add the header rows.
            for j in range(df.shape[-1]):
                t.cell(0, j).text = df.columns[j]
                # t.cell(0, j).fo

            # add the rest of the data frame
            for i in range(df.shape[0]):
                for j in range(df.shape[-1]):
                    t.cell(i + 1, j).text = str(df.values[i, j])

    document.save(pathfile)

def exportWord(p_vystup_docx):
    conn_str = 'oes_migr/oes_migr@oedev'
    db = OracleDatabase(conn_str)

    df = db.query('''
        select ut.TABLE_NAME, utc.COMMENTS, ut.TABLESPACE_NAME
          from user_tables ut
          join user_tab_comments utc
            on utc.TABLE_NAME = ut.TABLE_NAME
         where ut.TABLE_NAME like '%MIGRADO%'
         order by 1
    ''')

    doctable(df, 'TabulkaDM', p_vystup_docx)

    # # pro každou tabulku vytvoř Sheet
    # for table in df['TABLE_NAME']:
    #     print(table)
    #     df_table = db.query('''
    #         select utc.COLUMN_ID,
    #                utc.COLUMN_NAME,
    #                utc.DATA_TYPE,
    #                utc.NULLABLE,
    #                ucc.COMMENTS
    #           from user_tab_cols utc
    #           left join user_col_comments ucc
    #             on ucc.TABLE_NAME = utc.TABLE_NAME
    #            and ucc.COLUMN_NAME = utc.COLUMN_NAME
    #          where utc.TABLE_NAME = :table_name
    #          order by utc.COLUMN_ID
    #         ''', {'table_name': table})
    #
    #     df_table.to_excel(writer, sheet_name=table, index=False)
    #
    #     worksheet = writer.workbook().get_worksheet_by_name(table)
    #     for i, width in enumerate(get_col_widths(df_table)):
    #         worksheet.set_column(i, i, width)
    #
    # writer.save()

exportWord(r'c:\Users\root\Documents\CSOB\Migrado\test-MIGRADO_v3.docx')

